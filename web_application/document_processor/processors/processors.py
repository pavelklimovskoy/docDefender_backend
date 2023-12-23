import re
from docx import Document


class TxtProcessor:
    """
    Класс для обработки TXT
    """
    pass


class PNGProcessor:
    """
    Обработка png
    """
    pass


class JPGProcessor:
    """
    обработка jpf
    """
    pass


class PDFProcessor:
    """
    Обработка pdf
    """
    pass


class XMLProcessor:
    """
    Обработка xml
    """
    pass


class DocumentProcessorSelector:
    SUPPORTED_EXTENSION: list[str] = [".doc", "docx", ".txt", "png", "jpg", "pdf", "xml"]

    def _is_supported_extension(self, filename: str) -> bool:
        if filename[-3:] in self.SUPPORTED_EXTENSION:
            return True
        if filename[-4:] in self.SUPPORTED_EXTENSION:
            return True
        return False

    @classmethod
    def _get_file_extension(cls, file_name: str) -> str:
        separator_pos: int | None = None
        for i in range(len(file_name) - 1, 0, -1):
            if file_name[i] == '.':
                separator_pos = i
                break

        if separator_pos is None:
            raise Exception("Filename not contain extension")

        file_extension: str = file_name[separator_pos:]

        return file_extension

    @classmethod
    def get_processor(cls, filename: str, file_path: str):
        """
        Получение класса для обработки файла по его расширению
        :param filename:
        :param file_path:
        :return:
        """

        if not cls._is_supported_extension(filename):
            raise Exception("Unsupported filetype")

        match cls._get_file_extension(filename):
            case ".txt":
                return TxtProcessor
            case ".docx":
                return DocxProcessor
            case ".doc":
                return DocProcessor
            case ".png":
                return PNGProcessor
            case ".jpg":
                return JPGProcessor
            case ".pdf":
                return PDFProcessor
            case ".xml":
                return XMLProcessor
            case _:
                raise Exception("Unsupported filetype")


class DocProcessor:
    pass


class DocxProcessor:
    # PHONE_NUMBER_REGEXP = r'^((8|\+7)[\- ]?)?(\(?\d{3}\)?[\- ]?)?[\d\- ]{7,10}$'
    PHONE_NUMBER_REGEXP: str = r'\b\+?[7,8](\s*\d{3}\s*\d{3}\s*\d{2}\s*\d{2})\b'
    INN_REGEXP: str = r'ИНН'

    # def _get_str_length(self, word: str) -> int:
    #     return len(word)

    def _get_hided_number(self, number: str) -> str:
        return '*' * len(number)

    def _get_hided_inn(self):
        return '#' * 13

    def save_document(self):
        self.document.save(f'{self.output_document_name}.docx')

    def __init__(self, document_name: str):
        self.document_name: str = document_name
        self.output_document_name: str = f'{self.document_name}_anonymized'

        self.NUMBER_PATTERN = re.compile(self.PHONE_NUMBER_REGEXP)
        self.INN_PATTERN = re.compile(self.INN_REGEXP)

        self.document = Document(self.document_name)

    def stupid_phone_numbers_hiding(self):
        '''
        Простая замена всех вхождений номеров в документе
        '''

        for i, paragraph in enumerate(self.document.paragraphs):
            entries = self.NUMBER_PATTERN.findall(paragraph.text)

            processed_text = paragraph.text
            for j in range(1, len(processed_text) - 1):
                if processed_text[j].isdigit() and processed_text[j + 1] != '.' and processed_text[j - 1] != '.':
                    processed_text = processed_text.replace(processed_text[j], '#')

            paragraph.text = processed_text

        for i, table in enumerate(self.document.tables):
            for row in table.rows:
                for cell in row.cells:
                    processed_text = cell.text
                    cell.text = "*" * len(cell.text)
                    for ch in processed_text:
                        if ch.isdigit():
                            processed_text = processed_text.replace(ch, '#')

                    cell.text = processed_text

    def stupid_inn_hiding(self):
        '''
        Простая замена всех вхождений инн в документе
        '''

        for i, paragraph in enumerate(self.document.paragraphs):
            entries = self.INN_PATTERN.findall(paragraph.text)

            for entry in entries:
                print(entry)
                paragraph.text = paragraph.text.replace(
                    entry,
                    self._get_hided_number(self._get_hided_inn())
                )

    def _print_content(self) -> None:
        """
        Вывод содержимого документа в консоль
        :return:
        """
        for i, paragraph in enumerate(self.document.paragraphs):
            print(paragraph.text)

    def _is_contain_phone_numbers(self) -> bool:
        """
        Проверяем содержит ли документ телефонный номер
        :return:
        """
        res: bool = False

        return res

    def _is_contain_email(self) -> bool:
        """
        Проверяем содержит ли документ электронную почту
        :return:
        """
        pass

    def _is_contain_full_name(self) -> bool:
        """
        Проверяем наличие ФИО в документе
        :return:
        """
        pass


def main() -> None:
    doc = DocxProcessor(
        document_name="document.docx"
    )

    doc.stupid_phone_numbers_hiding()

    doc.stupid_inn_hiding()

    doc.save_document()


if __name__ == "__main__":
    # main()

    DocumentProcessorSelector._get_file_type("filename.txt")

'''
ОГРН 1026103165241,
ИНН 6163027810, КПП 616301001
'''
