import re
from docx import Document

from docDefender_backend import settings


class DocxProcessor:
    # PHONE_NUMBER_REGEXP = r'^((8|\+7)[\- ]?)?(\(?\d{3}\)?[\- ]?)?[\d\- ]{7,10}$'
    PHONE_NUMBER_REGEXP: str = r'\b\+?[7,8](\s*\d{3}\s*\d{3}\s*\d{2}\s*\d{2})\b'
    INN_REGEXP: str = r'ИНН'

    # def _get_str_length(self, word: str) -> int:
    #     return len(word)

    def _get_hided_number(self, number: str) -> str:
        return '*' * len(number)

    def _get_hided_inn(self):
        return '#' * 13

    def save_document(self):
        # self.document.save(f'{settings.MEDIA_ROOT}anon_{self.output_document_name}')
        self.document.save(f'{settings.MEDIA_ROOT}anon_{self.output_document_name}')

    def __init__(self, document_name: str):
        self.document_name: str = document_name
        self.output_document_name: str = f'{self.document_name}'

        self.NUMBER_PATTERN = re.compile(self.PHONE_NUMBER_REGEXP)
        self.INN_PATTERN = re.compile(self.INN_REGEXP)

        self.document = Document(settings.MEDIA_ROOT + self.document_name)

    def stupid_phone_numbers_hiding(self):
        '''
        Простая замена всех вхождений номеров в документе
        '''

        for i, paragraph in enumerate(self.document.paragraphs):
            entries = self.NUMBER_PATTERN.findall(paragraph.text)

            processed_text = paragraph.text
            for j in range(1, len(processed_text) - 1):
                if processed_text[j].isdigit() and processed_text[j + 1] != '.' and processed_text[j - 1] != '.':
                    processed_text = processed_text.replace(processed_text[j], '#')

            paragraph.text = processed_text

        for i, table in enumerate(self.document.tables):
            for row in table.rows:
                for cell in row.cells:
                    processed_text = cell.text
                    cell.text = "*" * len(cell.text)
                    for ch in processed_text:
                        if ch.isdigit():
                            processed_text = processed_text.replace(ch, '#')

                    cell.text = processed_text

    def stupid_inn_hiding(self):
        '''
        Простая замена всех вхождений инн в документе
        '''

        for i, paragraph in enumerate(self.document.paragraphs):
            entries = self.INN_PATTERN.findall(paragraph.text)

            for entry in entries:
                print(entry)
                paragraph.text = paragraph.text.replace(
                    entry,
                    self._get_hided_number(self._get_hided_inn())
                )

    def anonymize_doc(self) -> None:
        self.stupid_phone_numbers_hiding()
        self.stupid_inn_hiding()


    def _print_content(self) -> None:
        """
        Вывод содержимого документа в консоль
        :return:
        """
        for i, paragraph in enumerate(self.document.paragraphs):
            print(paragraph.text)

    def _is_contain_phone_numbers(self) -> bool:
        """
        Проверяем содержит ли документ телефонный номер
        :return:
        """
        res: bool = False

        return res

    def _is_contain_email(self) -> bool:
        """
        Проверяем содержит ли документ электронную почту
        :return:
        """
        pass

    def _is_contain_full_name(self) -> bool:
        """
        Проверяем наличие ФИО в документе
        :return:
        """
        pass

