import spacy
from docx import Document
from docDefender_backend import settings


class DocProcessor:

    def __init__(self, document_name: str):
        self.document_name: str = document_name
        self.output_document_name: str = f'{self.document_name}'

    def anonymize_doc(self):

        def anonymize_sensitive_info(text):
            nlp = spacy.load("en_core_web_sm")
            doc = nlp(text)

            for ent in doc.ents:
                if ent.label_ in ["PERSON", "ORG", "GPE", "DATE", "MONEY", "PHONE", "EMAIL", ]:
                    text = text.replace(ent.text, "*" * len(ent.text))

            return text

        
        document = Document(f'{settings.MEDIA_ROOT}{self.document_name}x')

        for paragraph in document.paragraphs:
            paragraph.text = anonymize_sensitive_info(paragraph.text)

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell.text = anonymize_sensitive_info(cell.text)

        document.save(f'{settings.MEDIA_ROOT}anon_{self.document_name}')

    def save_document(self) -> None:
        pass

    pass
